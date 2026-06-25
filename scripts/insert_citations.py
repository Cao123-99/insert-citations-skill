#!/usr/bin/env python3
"""
学术论文参考文献交叉引用插入工具 — 兼容 Word 和 WPS。

功能：
  1. 自动查找或创建「参考文献」节
  2. 解析参考文献条目（支持自动编号和手动编号）
  3. 按正文首次出现顺序重新排序参考文献
  4. 在参考文献段落上创建书签
  5. 将正文中的 [n] 替换为 REF 域（交叉引用）

用法：
  python insert_citations.py <docx路径> [--no-sort] [--sort-only] [--dry-run]
                          [--ref-style {plain,superscript,both}]

  --no-sort      跳过排序步骤，仅插入交叉引用字段
  --sort-only    只排序，不插入交叉引用
  --dry-run      仅分析和报告，不修改文档
  --ref-style    引用标记样式过滤：plain=仅匹配非上标, superscript=仅匹配上标,
                 both=匹配所有（默认）

需求：pip install python-docx
"""

import re
import sys
import os
import copy
import shutil
import random
import argparse
from dataclasses import dataclass, field
from typing import Optional
from collections import OrderedDict

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ═══════════════════════════════════════════════════════════
# 数据结构
# ═══════════════════════════════════════════════════════════

@dataclass
class Citation:
    """正文中的一处引用标记"""
    para_element: object = None       # lxml 元素（直接引用，避免通过 doc.paragraphs 索引）
    paragraph_index: int = 0          # 在 doc.paragraphs 中的索引（-1 表示表格单元格段落）
    position_in_text: int = 0
    original_numbers: list = field(default_factory=list)
    raw_text: str = ""
    is_field: bool = False
    ref_names: list = field(default_factory=list)
    is_superscript: bool = False     # True 表示该引用位于上标 run 中

@dataclass
class RefEntry:
    """一条参考文献"""
    paragraph_index: int
    original_number: int
    new_number: Optional[int] = None
    text: str = ""
    is_numbered_list: bool = True

# ═══════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════

def W_TAG(name):
    return f'{{{W}}}{name}'

def find_ref_heading(doc):
    """查找参考文献标题段落。返回 (段落索引, 段落) 或 (None, None)。"""
    for i, para in enumerate(doc.paragraphs):
        clean = re.sub(r'\s+', '', para.text)
        if re.search(r'参考文献', clean) or re.match(r'^[Rr]eferences?$', clean):
            return i, para
    return None, None

def has_auto_numbering(para):
    """段落是否使用 Word 自动编号（且不是标题）。"""
    pPr = para._element.find(W_TAG('pPr'))
    if pPr is not None:
        if pPr.find(W_TAG('outlineLvl')) is not None:
            return False
        if pPr.find(W_TAG('numPr')) is not None:
            return True
    return False

def get_reference_paragraphs(doc, heading_idx):
    """获取参考文献标题之后的所有参考文献条目段落索引。"""
    ref_paragraphs = []
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        clean = re.sub(r'\s+', '', text)

        # 遇到明显的新节标题 → 停止
        if re.match(r'^致\s*谢|^致谢$|^[Aa]cknowledge|^附录|^作者简介', clean):
            break

        if not text:
            continue

        is_ref = has_auto_numbering(para) or re.match(r'^\[\d+\]', text)
        # 排除短标签（图题/表题）
        if is_ref and len(text) < 30:
            if re.match(r'^(图|Fig|表|Table)\s*\d', clean):
                continue
        if is_ref:
            ref_paragraphs.append(i)

    return ref_paragraphs

def parse_ref_number(para, auto_index=None):
    """从参考文献段落提取编号。"""
    text = para.text.strip()
    m = re.match(r'^\s*\[(\d+)\]', text)
    if m:
        return int(m.group(1))
    if auto_index is not None and has_auto_numbering(para):
        return auto_index
    return None

def find_ref_fields(para_elem):
    """查找段落中已有的 REF 域。返回 [(ref_name, display_text), ...]"""
    results = []
    in_field = False
    after_sep = False
    current_ref = None
    display_parts = []

    for child in para_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            # 跳过书签等非 run 元素（它们可能穿插在 REF 域中）
            continue

        fld = child.find(W_TAG('fldChar'))
        instr = child.find(W_TAG('instrText'))
        t_elem = child.find(W_TAG('t'))

        if fld is not None:
            ft = fld.get(W_TAG('fldCharType'))
            if ft == 'begin':
                in_field = True; after_sep = False
                current_ref = None; display_parts = []
            elif ft == 'separate':
                after_sep = True
            elif ft == 'end':
                if current_ref:
                    results.append((current_ref, ''.join(display_parts)))
                in_field = False; after_sep = False
                current_ref = None; display_parts = []
            continue

        if in_field:
            if instr is not None:
                m = re.search(r'(_Ref\d+)', instr.text or '')
                if m:
                    current_ref = m.group(1)
            elif after_sep and t_elem is not None and t_elem.text:
                display_parts.append(t_elem.text)

    return results

def expand_number_string(s):
    """展开引用编号字符串。"1,3-5" → [1,3,4,5]。兼容 ][ 分隔（来自 [[17],[18]]）。"""
    s = s.replace('，', ',').replace('–', '-').replace('—', '-').replace(' ', '')
    s = s.replace('][', ',')  # [[17],[18]] → 17,18
    numbers = []
    for part in s.split(','):
        part = part.strip()
        if not part:
            continue
        if '-' in part:
            try:
                a, b = part.split('-', 1)
                numbers.extend(range(int(a), int(b) + 1))
            except ValueError:
                try:
                    numbers.append(int(part.replace('-', '')))
                except ValueError:
                    pass
        else:
            try:
                numbers.append(int(part))
            except ValueError:
                pass
    return numbers

def _build_superscript_ranges(para_elem):
    """
    构建段落中上标文本的字符位置区间列表。

    遍历段落 XML 中所有文本 run（跳过 REF 域的 fldChar/instrText），
    检查 rPr 中是否包含 vertAlign="superscript"，记录上标文本的 (start, end) 区间。

    Args:
        para_elem: <w:p> 段落的 lxml 元素

    Returns:
        [(start, end), ...] 列表，start 包含，end 不包含。无上标时返回空列表。
    """
    ranges = []
    char_pos = 0
    for child in para_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        if child.find(W_TAG('fldChar')) is not None:
            continue
        if child.find(W_TAG('instrText')) is not None:
            continue

        t_elem = child.find(W_TAG('t'))
        if t_elem is None or t_elem.text is None:
            continue

        rpr = child.find(W_TAG('rPr'))
        is_super = False
        if rpr is not None:
            va = rpr.find(W_TAG('vertAlign'))
            if va is not None and va.get(W_TAG('val')) == 'superscript':
                is_super = True

        text_len = len(t_elem.text)
        if is_super:
            ranges.append((char_pos, char_pos + text_len))
        char_pos += text_len

    return ranges


def parse_body_citations(doc, body_start, body_end, ref_style='both'):
    """从正文段落 + 表格单元格中提取所有引用标记。

    Args:
        doc: Document 对象
        body_start: 正文起始段落索引
        body_end: 正文结束段落索引
        ref_style: 匹配样式 — 'plain' 仅匹配非上标引用, 'superscript' 仅匹配上标引用,
                   'both' 匹配所有（默认）
    """
    citations = []

    def scan_paragraph(para, pi):
        """扫描单个段落中的引用"""
        para_text = para.text
        para_elem = para._element

        existing = find_ref_fields(para_elem)
        if existing:
            for ref_name, display in existing:
                citations.append(Citation(
                    para_element=para_elem,
                    paragraph_index=pi,
                    original_numbers=[int(ref_name.replace('_Ref', ''))],
                    raw_text=display,
                    is_field=True,
                    ref_names=[ref_name]
                ))
            return

        # 建立上标区间映射（用于样式过滤和格式记录）
        sup_ranges = _build_superscript_ranges(para_elem)

        def _is_position_superscript(pos):
            """检查字符位置是否落在上标区间内"""
            for start, end in sup_ranges:
                if start <= pos < end:
                    return True
            return False

        # 双方括号 [[n]] — 始终匹配，不受 ref_style 影响（格式足够明确）
        for m in re.finditer(r'\[\[([\d\s,\-，\–\—\[\]]+?)\]\]', para_text):
            nums = expand_number_string(m.group(1))
            nums = [n for n in nums if n > 0]
            if nums:
                citations.append(Citation(
                    para_element=para_elem,
                    paragraph_index=pi,
                    position_in_text=m.start(),
                    original_numbers=nums,
                    raw_text=m.group(0),
                    is_superscript=_is_position_superscript(m.start())
                ))

        # 单方括号 [n] — 按 ref_style 过滤
        for m in re.finditer(r'(?<!\[)\[([\d\s,\-，\–\—]+)\](?!\])', para_text):
            nums = expand_number_string(m.group(1))
            nums = [n for n in nums if n > 0]
            if not nums:
                continue

            match_is_superscript = _is_position_superscript(m.start())

            # 样式过滤
            if ref_style == 'plain' and match_is_superscript:
                continue   # plain 模式跳过上标引用
            if ref_style == 'superscript' and not match_is_superscript:
                continue   # superscript 模式跳过非上标引用

            citations.append(Citation(
                para_element=para_elem,
                paragraph_index=pi,
                position_in_text=m.start(),
                original_numbers=nums,
                raw_text=m.group(0),
                is_superscript=match_is_superscript
            ))

    # 1. 扫描正文段落
    for pi in range(body_start, body_end):
        scan_paragraph(doc.paragraphs[pi], pi)

    # 2. 扫描表格单元格中的段落（这些段落不在 doc.paragraphs 中）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        scan_paragraph(para, -1)  # -1 表示表格单元格

    return citations

def make_citation_text(numbers):
    """将数字列表压缩为引用文本 [1,3-5]"""
    numbers = sorted(set(numbers))
    parts = []
    i = 0
    while i < len(numbers):
        start = numbers[i]
        end = start
        while i + 1 < len(numbers) and numbers[i + 1] == end + 1:
            end = numbers[i + 1]; i += 1
        if end > start + 1:
            parts.append(f'{start}-{end}')
        elif end == start + 1:
            parts.append(f'{start},{end}')
        else:
            parts.append(str(start))
        i += 1
    return '[' + ','.join(parts) + ']'

def make_ref_field_runs(ref_number, base_rpr=None):
    """为单个引用编号创建 REF 域所需的 5 个 run 元素。
    显示文本的格式继承自 base_rpr（即正文段落格式），不再强制上标。"""
    runs = []

    # fldChar begin
    r = OxmlElement('w:r')
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    fc = OxmlElement('w:fldChar')
    fc.set(W_TAG('fldCharType'), 'begin')
    r.append(fc)
    runs.append(r)

    # instrText
    r = OxmlElement('w:r')
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    it = OxmlElement('w:instrText')
    it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    it.text = f' REF _Ref{ref_number} \\h '
    r.append(it)
    runs.append(r)

    # fldChar separate
    r = OxmlElement('w:r')
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    fc = OxmlElement('w:fldChar')
    fc.set(W_TAG('fldCharType'), 'separate')
    r.append(fc)
    runs.append(r)

    # result text — 显示 [n]（带方括号），格式继承段落正文样式，不上标
    r = OxmlElement('w:r')
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    t = OxmlElement('w:t')
    t.text = f'[{ref_number}]'
    r.append(t)
    runs.append(r)

    # fldChar end
    r = OxmlElement('w:r')
    if base_rpr is not None:
        r.append(copy.deepcopy(base_rpr))
    fc = OxmlElement('w:fldChar')
    fc.set(W_TAG('fldCharType'), 'end')
    r.append(fc)
    runs.append(r)

    return runs

# ═══════════════════════════════════════════════════════════
# 核心操作
# ═══════════════════════════════════════════════════════════

def build_renumbering_map(citations, ref_entries):
    """根据引用在正文中的首次出现顺序，建立旧编号→新编号映射。
    表格中的引用（paragraph_index < 0）排在所有正文段落引用之后。"""
    first_appearance = OrderedDict()
    for c in citations:
        for num in c.original_numbers:
            if num not in first_appearance:
                # 表格段落用 0-index 存储（负值），排序时修正
                pi = c.paragraph_index
                pos = c.position_in_text
                first_appearance[num] = (pi if pi >= 0 else 999999, pos)

    cited_order_raw = sorted(first_appearance.keys(), key=lambda n: first_appearance[n])
    all_refs = {r.original_number for r in ref_entries if r.original_number}
    # 只保留参考文献列表中实际存在的编号（过滤图表标签等误识别）
    cited_order = [n for n in cited_order_raw if n in all_refs]
    uncited = all_refs - set(cited_order)

    mapping = {}
    for new_num, old_num in enumerate(cited_order + sorted(uncited), 1):
        mapping[old_num] = new_num
    return mapping

def ensure_auto_numbering(doc, elements):
    """
    确保参考文献段落使用 Word 自动编号（[1] [2] ... 格式）。
    如已有有效编号定义则直接复用；否则在 numbering_part 中创建新定义。
    """
    try:
        num_part = doc.part.numbering_part
    except Exception:
        print("    ⚠️ 无法访问 numbering_part，使用手动 [n] 文本前缀")
        return None

    if num_part is None:
        print("    ⚠️ numbering_part 不存在，使用手动 [n] 文本前缀")
        return None

    numbering = num_part.element

    # 1. 检查段落是否已有有效 numPr（numId 在 numbering 中有定义）
    existing_numIds = set()
    for elem in elements:
        pPr = elem.find(W_TAG('pPr'))
        if pPr is not None:
            np = pPr.find(W_TAG('numPr'))
            if np is not None:
                ni = np.find(W_TAG('numId'))
                if ni is not None:
                    existing_numIds.add(ni.get(W_TAG('val')))

    defined_ids = set()
    for num_el in numbering.findall(W_TAG('num')):
        defined_ids.add(num_el.get(W_TAG('numId')))

    if existing_numIds and existing_numIds.issubset(defined_ids):
        # 已有有效编号 — 无需创建
        return list(existing_numIds)[0]

    # 2. 如果段落有 numPr 但编号不存在，清除这些坏的 numPr
    for elem in elements:
        pPr = elem.find(W_TAG('pPr'))
        if pPr is not None:
            np = pPr.find(W_TAG('numPr'))
            if np is not None:
                pPr.remove(np)

    # 3. 创建新的编号定义
    print(f"   为参考文献添加 Word 自动编号 ([1] [2] ... 格式)")

    all_ids = set()
    for n in numbering.findall(W_TAG('num')):
        all_ids.add(n.get(W_TAG('numId')))
    for an in numbering.findall(W_TAG('abstractNum')):
        all_ids.add(an.get(W_TAG('abstractNumId')))

    new_id = '99'
    while new_id in all_ids:
        new_id = str(int(new_id) + 1)

    nsid_val = ''.join(random.choice('0123456789ABCDEF') for _ in range(8))

    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(W_TAG('abstractNumId'), new_id)

    nsid_el = OxmlElement('w:nsid')
    nsid_el.set(W_TAG('val'), nsid_val)
    abstractNum.append(nsid_el)

    mt = OxmlElement('w:multiLevelType')
    mt.set(W_TAG('val'), 'hybridMultilevel')
    abstractNum.append(mt)

    lvl = OxmlElement('w:lvl')
    lvl.set(W_TAG('ilvl'), '0')

    start = OxmlElement('w:start')
    start.set(W_TAG('val'), '1')
    lvl.append(start)

    numFmt = OxmlElement('w:numFmt')
    numFmt.set(W_TAG('val'), 'decimal')
    lvl.append(numFmt)

    lvlText = OxmlElement('w:lvlText')
    lvlText.set(W_TAG('val'), '[%1]')
    lvl.append(lvlText)

    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(W_TAG('val'), 'left')
    lvl.append(lvlJc)

    lvl_pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(W_TAG('left'), '420')
    ind.set(W_TAG('hanging'), '420')
    lvl_pPr.append(ind)
    lvl.append(lvl_pPr)

    abstractNum.append(lvl)
    numbering.append(abstractNum)

    num_el = OxmlElement('w:num')
    num_el.set(W_TAG('numId'), new_id)
    aid_ref = OxmlElement('w:abstractNumId')
    aid_ref.set(W_TAG('val'), new_id)
    num_el.append(aid_ref)
    numbering.append(num_el)

    for elem in elements:
        pPr = elem.find(W_TAG('pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            elem.insert(0, pPr)

        numPr_new = OxmlElement('w:numPr')
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(W_TAG('val'), new_id)
        numPr_new.append(numId_elem)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(W_TAG('val'), '0')
        numPr_new.append(ilvl)
        pPr.append(numPr_new)

    return new_id

def reorder_references_v2(doc, ref_entries, mapping):
    """
    重新排序参考文献。

    保留自动编号，物理重排段落元素。重排后自动编号自动反映新顺序。
    返回 (sorted_entries, sorted_elements)。
    """
    for entry in ref_entries:
        entry.new_number = mapping.get(entry.original_number, entry.original_number)

    sorted_entries = sorted(ref_entries, key=lambda e: e.new_number or 9999)

    # 收集所有参考文献段落的 lxml 元素
    # 注意：sorted_entries 按 new_number 排序，元素收集也按此顺序
    elements_to_reorder = []
    for entry in sorted_entries:
        para = doc.paragraphs[entry.paragraph_index]
        elements_to_reorder.append(para._element)

    if len(elements_to_reorder) <= 1:
        return sorted_entries, elements_to_reorder

    body = doc.element.body

    # ★ 关键：锚点必须是引用块之前、不属于引用块的元素（即参考文献标题）
    # elements_to_reorder 按 new_number 排序，其第一个元素不一定是文档中的第一个引用段落
    # 因此找所有引用段落中 paragraph_index 最小的那个（文档中最早出现的引用段落）
    min_pi = min(e.paragraph_index for e in sorted_entries)
    anchor = doc.paragraphs[min_pi]._element.getprevious()  # 标题段落

    if anchor is None:
        print("   ⚠️  锚点为 None——参考文献段落在文档开头？跳过重排")
        return sorted_entries, elements_to_reorder

    # 验证锚点
    anchor_text = ''
    for t in anchor.iter(W_TAG('t')):
        if t.text:
            anchor_text += t.text
    print(f"   锚点: {anchor_text[:60].strip()}...")

    # 从树中移除所有参考文献段落
    for elem in elements_to_reorder:
        body.remove(elem)

    # 按新顺序在锚点之后重新插入
    insert_after = anchor
    for elem in elements_to_reorder:
        insert_after.addnext(elem)
        insert_after = elem

    # 确保参考文献使用 Word 自动编号（[1] [2] ... 格式）
    num_id = ensure_auto_numbering(doc, elements_to_reorder)

    # 清理所有残留的手动 [n] 文本前缀（避免与自动编号重复）
    for entry, elem in zip(sorted_entries, elements_to_reorder):
        text_runs = []
        for run in elem.findall(W_TAG('r')):
            t = run.find(W_TAG('t'))
            if t is not None and t.text is not None:
                text_runs.append((run, t))
        if text_runs:
            full_text = ''.join(t.text for _, t in text_runs)
            # 去掉所有已有的 [n] 前缀
            full_text = re.sub(r'^\s*\[\d+\]\s*', '', full_text)
            for run, t in text_runs:
                t.text = ''
            text_runs[0][1].text = full_text

    print(f"   已重排 {len(elements_to_reorder)} 个段落（保留自动编号）")

    return sorted_entries, elements_to_reorder

def update_body_ref_targets(doc, mapping):
    """扫描文档中所有 REF 域，将其 instrText 中的 _Ref<n> 目标更新为新编号。"""
    count = 0
    for para in doc.paragraphs:
        for child in para._element.iter():
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'instrText' and child.text and 'REF _Ref' in child.text:
                m = re.search(r'REF (_Ref)(\d+)', child.text)
                if m:
                    old_num = int(m.group(2))
                    if old_num in mapping:
                        new_num = mapping[old_num]
                        child.text = child.text.replace(
                            f'{m.group(1)}{old_num}',
                            f'{m.group(1)}{new_num}'
                        )
                        count += 1
    return count

def update_body_display_texts(doc, mapping):
    """
    更新 REF 域的显示文本 [n] → [new_n]。
    同时移除 result run 上的上标格式，改为与正文一致的格式。
    """
    count = 0
    in_field = False
    after_sep = False

    for para in doc.paragraphs:
        for child in para._element.iter():
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'fldChar':
                ft = child.get(W_TAG('fldCharType'))
                if ft == 'begin':
                    in_field = True; after_sep = False
                elif ft == 'separate':
                    after_sep = True
                elif ft == 'end':
                    in_field = False; after_sep = False
                continue

            if tag == 't' and child.text and child.text.strip() and in_field and after_sep:
                stripped = child.text.strip()
                for old_num, new_num in mapping.items():
                    # 各种显示格式： [n] 或 n（旧版纯数字）
                    if stripped == f'[{old_num}]' or stripped == str(old_num):
                        child.text = f'[{new_num}]'
                        count += 1
                        # 剥离上标格式，统一用正文格式
                        parent_run = child.getparent()
                        if parent_run is not None:
                            run_rpr = parent_run.find(W_TAG('rPr'))
                            if run_rpr is not None:
                                va = run_rpr.find(W_TAG('vertAlign'))
                                if va is not None and va.get(W_TAG('val')) == 'superscript':
                                    run_rpr.remove(va)
                                sz = run_rpr.find(W_TAG('sz'))
                                if sz is not None and sz.get(W_TAG('val')) == '18':
                                    run_rpr.remove(sz)
                                if len(run_rpr) == 0:
                                    parent_run.remove(run_rpr)
                        break
    return count

def insert_ref_fields_for_plain_text(doc, citations):
    """将正文中纯文本的 [n] / [[n]] 引用替换为 REF 域。"""
    for c in reversed(citations):
        if c.is_field:
            continue

        para_elem = c.para_element

        # 根据引用是否上标选择对应的 rPr，保留原始格式
        if c.is_superscript:
            base_rpr = _get_superscript_rpr(para_elem)
            if base_rpr is None:
                # 回退：段落中没有上标 rPr，使用正文格式
                base_rpr = _get_body_rpr(para_elem)
        else:
            base_rpr = _get_body_rpr(para_elem)

        # REF 域自带 [n] 显示，多引用时用逗号分隔
        nums = sorted(set(c.original_numbers))
        field_runs = []

        for i, num in enumerate(nums):
            if i > 0:
                r = OxmlElement('w:r')
                if base_rpr is not None:
                    r.append(copy.deepcopy(base_rpr))
                t = OxmlElement('w:t')
                t.text = ','
                r.append(t)
                field_runs.append(r)
            field_runs.extend(make_ref_field_runs(num, base_rpr))

        # 在段落中定位并替换
        _replace_text_with_field_runs(para_elem, c.raw_text, c.position_in_text, field_runs)


def _get_body_rpr(para_elem):
    """
    从段落中取正文 run 的 rPr（跳过 REF 域、上标 run）。
    如果段落没有合适的正文 run，返回 None（使用 Word 默认格式）。
    """
    best = None
    for child in para_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        if child.find(W_TAG('fldChar')) is not None:
            continue
        if child.find(W_TAG('instrText')) is not None:
            continue

        t_elem = child.find(W_TAG('t'))
        if t_elem is None or not t_elem.text or not t_elem.text.strip():
            continue

        rpr = child.find(W_TAG('rPr'))
        if rpr is None:
            # 无 rPr 的正文 run —— 最"干净"的格式，直接用
            return None

        va = rpr.find(W_TAG('vertAlign'))
        if va is not None and va.get(W_TAG('val')) == 'superscript':
            # 跳过上标 run，取其格式但不带上标属性
            rpr_copy = copy.deepcopy(rpr)
            rpr_copy.remove(va)
            # 同时移除小字号（上标通常 sz=18）
            sz = rpr_copy.find(W_TAG('sz'))
            if sz is not None and sz.get(W_TAG('val')) == '18':
                rpr_copy.remove(sz)
            if len(rpr_copy) > 0:
                best = rpr_copy
            continue

        # 普通正文 run —— 最优先
        if best is None:
            best = copy.deepcopy(rpr)
            break

    return best


def _get_superscript_rpr(para_elem):
    """
    从段落中获取上标 run 的 rPr（含 vertAlign="superscript"）。

    返回第一个上标文本 run 的 rPr 深拷贝，用于创建保持上标格式的 REF 域。
    若段落中没有上标 run，返回 None。

    Args:
        para_elem: <w:p> 段落的 lxml 元素

    Returns:
        rPr 元素的深拷贝，或 None
    """
    for child in para_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        if child.find(W_TAG('fldChar')) is not None:
            continue
        if child.find(W_TAG('instrText')) is not None:
            continue

        rpr = child.find(W_TAG('rPr'))
        if rpr is None:
            continue

        va = rpr.find(W_TAG('vertAlign'))
        if va is not None and va.get(W_TAG('val')) == 'superscript':
            return copy.deepcopy(rpr)

    return None


def _replace_text_with_field_runs(para_elem, old_text, position, field_runs):
    """在段落 XML 中定位纯文本引用并替换为 REF 域 runs。
    支持引用文本跨多个 run 的情况（如 [ 、4 、] 分别在三个 run 中）。"""
    # 收集所有纯文本 run（跳过 REF 域）
    text_runs = []
    for child in list(para_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        if child.find(W_TAG('fldChar')) is not None:
            continue
        if child.find(W_TAG('instrText')) is not None:
            continue
        t_elem = child.find(W_TAG('t'))
        if t_elem is not None and t_elem.text is not None:
            text_runs.append(child)

    # 拼接文本定位
    full_text = ''.join(run.find(W_TAG('t')).text for run in text_runs)

    if position + len(old_text) > len(full_text):
        # 位置偏移，尝试模糊匹配
        idx = full_text.find(old_text, max(0, position - 5))
        if idx >= 0:
            position = idx
        else:
            return

    # 找到覆盖 old_text 的所有 run
    char_pos = 0
    involved_runs = []
    for run in text_runs:
        t = run.find(W_TAG('t'))
        run_start = char_pos
        run_end = char_pos + len(t.text)
        char_pos = run_end

        if run_end <= position or run_start >= position + len(old_text):
            continue  # 不涉及
        involved_runs.append(run)

    if not involved_runs:
        return

    # 处理涉及的 run：切分、替换
    first_run = involved_runs[0]
    last_run = involved_runs[-1]

    first_t = first_run.find(W_TAG('t'))
    last_t = last_run.find(W_TAG('t'))

    # 计算在第一个 run 中的截断位置
    first_start = position - sum(len(r.find(W_TAG('t')).text) for r in text_runs[:text_runs.index(first_run)])
    first_end = min(len(first_t.text), first_start + len(old_text))

    # 保留第一个 run 的前半段
    before = first_t.text[:first_start]

    # 保留最后一个 run 的后半段
    last_pos_in_text = position + len(old_text) - sum(len(r.find(W_TAG('t')).text) for r in text_runs[:text_runs.index(last_run)])
    after = last_t.text[last_pos_in_text:]

    # 删除涉及的中间 run（除了 first_run 和 last_run）
    middle_runs = involved_runs[1:-1]  # first 和 last 之间
    for run in middle_runs:
        para_elem.remove(run)

    # 如果跨多个 run，删除 last_run（现在只剩 first 和 last 之间没有其他 run 阻隔）
    if len(involved_runs) > 1:
        para_elem.remove(last_run)

    # 更新第一个 run 的文本
    first_t.text = before

    # 在 first_run 之后插入 REF 域 runs
    insert_after = first_run
    for fr in field_runs:
        insert_after.addnext(fr)
        insert_after = fr

    # 插入后段文本
    if after:
        r = OxmlElement('w:r')
        old_rpr = last_run.find(W_TAG('rPr'))
        if old_rpr is not None:
            r.append(copy.deepcopy(old_rpr))
        t = OxmlElement('w:t')
        t.text = after
        r.append(t)
        insert_after.addnext(r)

def delete_and_create_bookmarks(doc, sorted_entries, sorted_elements):
    """删除旧 _Ref 书签，在重新排序后的段落上创建新的 _Ref1.._RefN 书签。"""
    body = doc.element.body

    # 删除旧的 _Ref bookmarks
    for bk in list(body.iter(W_TAG('bookmarkStart'))):
        name = bk.get(W_TAG('name'))
        if name and re.match(r'^_Ref\d+$', name):
            parent = bk.getparent()
            if parent is not None:
                parent.remove(bk)

    # 创建新书签
    for entry, elem in zip(sorted_entries, sorted_elements):
        bookmark_name = f'_Ref{entry.new_number}'
        bookmark_id = 60000 + (entry.new_number or 0)

        first_run = elem.find(W_TAG('r'))
        if first_run is None:
            first_run = OxmlElement('w:r')
            elem.insert(0, first_run)

        bk_start = OxmlElement('w:bookmarkStart')
        bk_start.set(W_TAG('id'), str(bookmark_id))
        bk_start.set(W_TAG('name'), bookmark_name)

        bk_end = OxmlElement('w:bookmarkEnd')
        bk_end.set(W_TAG('id'), str(bookmark_id))

        first_run.addprevious(bk_start)
        elem.append(bk_end)

def verify_result(doc):
    """验证修复结果。"""
    body = doc.element.body
    bookmarks = set()
    ref_fields = set()

    for bk in body.iter(W_TAG('bookmarkStart')):
        name = bk.get(W_TAG('name'))
        if name and re.match(r'^_Ref\d+$', name):
            bookmarks.add(name)

    for para in doc.paragraphs:
        for m in re.finditer(r'REF (_Ref\d+)', para._element.xml):
            ref_fields.add(m.group(1))

    print(f"\n🔍 验证结果：")
    print(f"   书签数量：{len(bookmarks)}")
    print(f"   REF 域数量：{len(ref_fields)}")

    missing = ref_fields - bookmarks
    if missing:
        print(f"   ❌ 缺失书签：{sorted(missing, key=lambda x: int(x.replace('_Ref','')))}")
    else:
        print(f"   ✅ 所有 REF 域都有对应的书签")

# ═══════════════════════════════════════════════════════════
# 报告
# ═══════════════════════════════════════════════════════════

def report_analysis(doc, citations, ref_entries, heading_idx):
    """打印分析报告。"""
    print("=" * 60)
    print("📊 文档分析报告")
    print("=" * 60)

    if heading_idx is not None:
        print(f"\n✅ 找到参考文献标题：P{heading_idx}")
        print(f"   标题：{doc.paragraphs[heading_idx].text.strip()[:80]}")
    else:
        print(f"\n⚠️  未找到参考文献标题")

    print(f"\n📚 参考文献：{len(ref_entries)} 条")
    for e in ref_entries[:5]:
        print(f"   [{e.original_number}] {e.text[:80]}...")
    if len(ref_entries) > 5:
        print(f"   ... 共 {len(ref_entries)} 条")

    print(f"\n📎 引用标记：{len(citations)} 处")
    for c in citations[:10]:
        tag = "(REF域)" if c.is_field else ""
        print(f"   P{c.paragraph_index}: {c.raw_text} {tag}")
    if len(citations) > 10:
        print(f"   ... 共 {len(citations)} 处")

    all_cited = set()
    for c in citations:
        all_cited.update(c.original_numbers)
    all_refs = {e.original_number for e in ref_entries}
    cited = all_refs & all_cited
    uncited = all_refs - all_cited
    undefined = all_cited - all_refs

    # 引用覆盖统计
    total_refs = len(all_refs)
    cited_count = len(cited)
    uncited_count = len(uncited)
    coverage = cited_count / total_refs * 100 if total_refs > 0 else 0

    print(f"\n{'─' * 58}")
    print(f"📊 引用覆盖统计")
    print(f"{'─' * 58}")
    print(f"  参考文献总数：{total_refs} 条")
    print(f"  已被引用：{cited_count} 条 ({coverage:.0f}%)")
    print(f"  未被引用：{uncited_count} 条 ({100-coverage:.0f}%)")

    if uncited:
        # 用紧凑格式显示编号（如 [4,5,12,13,30-38]）
        uncited_list = sorted(uncited)
        compact = []
        i = 0
        while i < len(uncited_list):
            start = uncited_list[i]
            end = start
            while i + 1 < len(uncited_list) and uncited_list[i + 1] == end + 1:
                end = uncited_list[i + 1]; i += 1
            if end > start + 1:
                compact.append(f'{start}-{end}')
            elif end == start + 1:
                compact.append(f'{start},{end}')
            else:
                compact.append(str(start))
            i += 1
        print(f"  未引用编号：[{','.join(compact)}]")

    if undefined:
        undefined_list = sorted(undefined)
        print(f"\n❌ 引用但未定义：{len(undefined_list)} 处编号 [{','.join(map(str, undefined_list))}]")
    else:
        print(f"{'─' * 58}")
        if coverage == 100:
            print(f"  ✅ 全部参考文献均被引用")
        elif coverage == 0:
            print(f"  ❌ 没有任何参考文献被引用——可能正文引用标记格式不标准")
        else:
            print(f"  ⚠️  {uncited_count} 条文献未被引用，请确认是否需要保留")
        print(f"{'─' * 58}")

# ═══════════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════════

def process_document(input_path, output_path=None, do_sort=True, sort_only=False,
                     dry_run=False, ref_style='both'):
    """处理文档的主流程。"""
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_cited{ext}"

    doc = Document(input_path)

    # ── 步骤 1：定位参考文献节 ──
    print("\n📋 步骤 1：定位参考文献节")
    heading_idx, heading_para = find_ref_heading(doc)

    if heading_idx is None:
        print('   未找到"参考文献"标题，将在文档末尾创建')
        if not dry_run:
            _create_ref_heading(doc)
            doc.save(output_path)
            doc = Document(output_path)
            heading_idx, heading_para = find_ref_heading(doc)
            if heading_idx is not None:
                print(f'   ✅ 已在 P{heading_idx} 创建"参考文献"标题')

    if heading_idx is None:
        print("   ❌ 无法创建/定位参考文献标题"); return

    # ── 步骤 2：解析参考文献条目 ──
    print("\n📋 步骤 2：解析参考文献条目")
    ref_para_indices = get_reference_paragraphs(doc, heading_idx)

    ref_entries = []
    auto_counter = 1
    for rpi in ref_para_indices:
        para = doc.paragraphs[rpi]
        text = para.text.strip()
        if not text:
            continue

        is_auto = has_auto_numbering(para)
        if is_auto:
            ref_num = parse_ref_number(para, auto_index=auto_counter)
            auto_counter += 1
        else:
            ref_num = parse_ref_number(para)
            if ref_num:
                auto_counter = ref_num + 1

        if ref_num is None:
            continue

        ref_entries.append(RefEntry(
            paragraph_index=rpi,
            original_number=ref_num,
            text=text,
            is_numbered_list=is_auto
        ))

    if not ref_entries:
        print("   ❌ 未找到参考文献条目")
        return
    print(f"   ✅ 找到 {len(ref_entries)} 条参考文献")

    # ── 步骤 3：解析引用标记 ──
    print("\n📋 步骤 3：解析正文中的引用标记")

    # 正文范围：摘要/关键词之后，参考文献之前
    body_start = 0
    for i, para in enumerate(doc.paragraphs):
        clean = re.sub(r'\s+', '', para.text)
        if clean in ('摘要', 'Abstract') or clean.startswith('摘要') or clean.startswith('Abstract'):
            body_start = i + 1
        if clean.startswith('关键词') or clean.startswith('Keywords'):
            body_start = i + 1

    # 跳到第一个长段落（正文开始）
    for i in range(body_start, min(heading_idx, len(doc.paragraphs))):
        if len(doc.paragraphs[i].text.strip()) > 50:
            body_start = i
            break

    body_end = heading_idx
    citations = parse_body_citations(doc, body_start, body_end, ref_style=ref_style)
    print(f"   ✅ 找到 {len(citations)} 处引用标记")

    report_analysis(doc, citations, ref_entries, heading_idx)

    if dry_run:
        print("\n🔍 --dry-run 模式，不修改文档")
        return

    # ── 步骤 4：排序 ──
    sorted_elements = None
    if do_sort:
        print("\n📋 步骤 4：按首次出现顺序排序参考文献")
        mapping = build_renumbering_map(citations, ref_entries)

        if mapping:
            print(f"   重编号映射：{len(mapping)} 条")
            # 更新正文 REF 域
            n1 = update_body_ref_targets(doc, mapping)
            n2 = update_body_display_texts(doc, mapping)
            print(f"   已更新 {n1} 个 REF 域目标 + {n2} 个显示文本")

            # 更新纯文本引用编号，同时更新 citation 对象
            for c in reversed(citations):
                if c.is_field:
                    continue
                new_nums = sorted(set(mapping.get(n, n) for n in c.original_numbers))
                new_text = make_citation_text(new_nums)
                # 双方括号 [[n]] → 统一转为 [n]（后续步骤 6 会插入 REF 域）
                if new_text != c.raw_text:
                    _replace_text_in_elem(c.para_element,
                                         c.raw_text, new_text, c.position_in_text)
                    c.raw_text = new_text
                    c.original_numbers = new_nums

            # 重排参考文献段落
            sorted_entries, sorted_elements = reorder_references_v2(doc, ref_entries, mapping)
            ref_entries = sorted_entries
            print(f"   ✅ 排序完成")
        else:
            print(f"   ⚠️  无需排序")
            sorted_elements = [doc.paragraphs[e.paragraph_index]._element for e in ref_entries]
    else:
        # no-sort: 确保 new_number 已设置
        for e in ref_entries:
            e.new_number = e.original_number
        sorted_elements = [doc.paragraphs[e.paragraph_index]._element for e in ref_entries]

    # ── 步骤 5：创建书签 ──
    if not sort_only:
        print("\n📋 步骤 5：在参考文献上创建书签")
        delete_and_create_bookmarks(doc, ref_entries, sorted_elements)
        print(f"   ✅ 创建了 {len(ref_entries)} 个书签")

    # ── 步骤 6：插入 REF 域 ──
    if not sort_only:
        print("\n📋 步骤 6：将纯文本引用替换为 REF 域")
        fresh_citations = parse_body_citations(doc, body_start, body_end, ref_style=ref_style)
        valid_ref_nums = {e.new_number or e.original_number for e in ref_entries}
        plain_citations = [c for c in fresh_citations
                           if not c.is_field
                           and all(n in valid_ref_nums for n in c.original_numbers)]
        skipped = len([c for c in fresh_citations if not c.is_field]) - len(plain_citations)
        insert_ref_fields_for_plain_text(doc, plain_citations)
        if skipped:
            print(f"   ⚠️  跳过 {skipped} 处无效引用（编号不在参考文献列表中）")
        print(f"   ✅ 已处理 {len(plain_citations)} 处引用")
    else:
        print("\n📋 --sort-only 模式：已排序，跳过交叉引用")

    # ── 保存 ──
    doc.save(output_path)
    print(f"\n{'=' * 60}")
    print(f"✅ 已保存：{output_path}")

    verify_doc = Document(output_path)
    verify_result(verify_doc)

    print(f"\n💡 在 Word/WPS 中打开后按 Ctrl+A → F9 更新所有域")

def _create_ref_heading(doc):
    """在文档末尾创建「参考文献」标题。"""
    body = doc.element.body
    # 找到最后一个段落之后的位置
    last_section = None
    for child in body:
        if child.tag == W_TAG('sectPr'):
            last_section = child
            break

    heading = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(W_TAG('val'), 'center')
    pPr.append(jc)
    heading.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(W_TAG('val'), '28')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = '参考文献'
    r.append(t)
    heading.append(r)

    if last_section is not None:
        last_section.addprevious(heading)
    else:
        body.append(heading)

def _replace_text_in_elem(para_elem, old_text, new_text, position):
    """在段落 XML 元素中替换指定位置的字符串。"""
    char_pos = 0
    for run in para_elem.findall(W_TAG('r')):
        t = run.find(W_TAG('t'))
        if t is None or t.text is None:
            continue
        if char_pos <= position < char_pos + len(t.text):
            local = position - char_pos
            if t.text[local:local + len(old_text)] == old_text:
                t.text = t.text[:local] + new_text + t.text[local + len(old_text):]
            return
        char_pos += len(t.text)

# ═══════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description='学术论文参考文献交叉引用插入工具')
    parser.add_argument('input', help='输入 docx 文件路径')
    parser.add_argument('-o', '--output', help='输出路径（默认：原名_cited.docx）')
    parser.add_argument('--no-sort', action='store_true', help='只插入交叉引用，不排序')
    parser.add_argument('--sort-only', action='store_true', help='只排序，不插入交叉引用')
    parser.add_argument('--dry-run', action='store_true', help='仅分析不修改')
    parser.add_argument('--ref-style', choices=['plain', 'superscript', 'both'],
                        default='both',
                        help='引用标记样式过滤：plain=仅匹配非上标, '
                             'superscript=仅匹配上标, both=匹配所有（默认）')

    args = parser.parse_args()
    if not os.path.exists(args.input):
        print(f"❌ 文件不存在：{args.input}")
        sys.exit(1)

    process_document(args.input, args.output, do_sort=not args.no_sort,
                     sort_only=args.sort_only, dry_run=args.dry_run,
                     ref_style=args.ref_style)

if __name__ == '__main__':
    main()
